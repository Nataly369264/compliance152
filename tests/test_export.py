"""Tests for DOCX export, PDF export, and API endpoints."""
from __future__ import annotations

import io

import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

from src.export.docx_converter import MarkdownToDocxConverter
from src.export.pdf_converter import MarkdownToPdfConverter, create_merged_pdf


# ── Heading tests ────────────────────────────────────────────────


def test_convert_heading1():
    """H1 heading produces a Heading 1 styled paragraph."""
    converter = MarkdownToDocxConverter(title="Test")
    doc = converter.convert("# Main Title")
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert len(headings) == 1
    assert headings[0].text == "Main Title"


def test_convert_heading2():
    """H2 heading produces a Heading 2 styled paragraph."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("## Section")
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    assert len(headings) == 1
    assert headings[0].text == "Section"


def test_convert_heading3():
    """H3 heading produces a Heading 3 styled paragraph."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("### Subsection")
    headings = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert len(headings) == 1
    assert headings[0].text == "Subsection"


def test_convert_multiple_headings():
    """Multiple heading levels are rendered correctly."""
    md = "# Title\n\n## Section 1\n\n### Sub 1.1\n\n## Section 2"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    h1 = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    h2 = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    h3 = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert len(h1) == 1
    assert len(h2) == 2
    assert len(h3) == 1


# ── Paragraph tests ──────────────────────────────────────────────


def test_convert_plain_paragraph():
    """Regular text becomes a Normal-styled paragraph."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("This is a plain paragraph.")
    normal = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    assert len(normal) >= 1
    assert "This is a plain paragraph." in normal[0].text


def test_convert_multiline_paragraph():
    """Consecutive non-blank lines are joined into a single paragraph."""
    md = "Line one\nline two\nline three"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    normal = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    assert len(normal) == 1
    assert "Line one" in normal[0].text
    assert "line three" in normal[0].text


def test_convert_two_paragraphs_separated_by_blank():
    """Two paragraphs separated by blank line produce two Normal paragraphs."""
    md = "First paragraph.\n\nSecond paragraph."
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    normal = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    assert len(normal) == 2


# ── List tests ───────────────────────────────────────────────────


def test_convert_bullet_list():
    """Bullet list items are converted to List Bullet paragraphs."""
    md = "- Item one\n- Item two\n- Item three"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullets) == 3
    assert "Item one" in bullets[0].text
    assert "Item three" in bullets[2].text


def test_convert_bullet_list_with_asterisk():
    """Asterisk bullet items also work."""
    md = "* Alpha\n* Beta"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullets) == 2


def test_convert_numbered_list():
    """Numbered list items contain the number prefix."""
    md = "1. First item\n2. Second item\n3. Third item"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    normal = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()]
    assert len(normal) == 3
    assert "1." in normal[0].text
    assert "First item" in normal[0].text


# ── Formatting tests ─────────────────────────────────────────────


def test_convert_bold_text():
    """Bold **text** creates a bold run."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("This is **bold** text")
    para = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()][0]
    bold_runs = [r for r in para.runs if r.bold]
    assert len(bold_runs) >= 1
    assert any("bold" in r.text for r in bold_runs)


def test_convert_italic_text():
    """Italic *text* creates an italic run."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("This is *italic* text")
    para = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()][0]
    italic_runs = [r for r in para.runs if r.italic]
    assert len(italic_runs) >= 1
    assert any("italic" in r.text for r in italic_runs)


def test_convert_mixed_formatting():
    """Paragraph with both bold and italic runs."""
    md = "Text with **bold** and *italic* words."
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    para = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text.strip()][0]
    bold_runs = [r for r in para.runs if r.bold]
    italic_runs = [r for r in para.runs if r.italic]
    assert len(bold_runs) >= 1
    assert len(italic_runs) >= 1


# ── Horizontal rule ──────────────────────────────────────────────


def test_convert_horizontal_rule():
    """--- creates a paragraph (horizontal rule separator)."""
    md = "Before\n\n---\n\nAfter"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    # Should have at least 3 paragraphs: before, rule, after
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Before" in texts
    assert "After" in texts


# ── Table tests ──────────────────────────────────────────────────


def test_convert_table():
    """Pipe-delimited table produces a DOCX table."""
    md = "| Col A | Col B |\n|-------|-------|\n| val1 | val2 |"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Col A"
    assert table.rows[0].cells[1].text == "Col B"
    assert table.rows[1].cells[0].text == "val1"
    assert table.rows[1].cells[1].text == "val2"


def test_convert_table_header_bold():
    """Table header row should be bold."""
    md = "| Header 1 | Header 2 |\n|----------|----------|\n| data 1 | data 2 |"
    converter = MarkdownToDocxConverter()
    doc = converter.convert(md)
    header_cell = doc.tables[0].rows[0].cells[0]
    assert any(r.bold for r in header_cell.paragraphs[0].runs)


# ── Style tests ──────────────────────────────────────────────────


def test_document_uses_times_new_roman():
    """Normal style font is Times New Roman."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("Test text")
    assert doc.styles["Normal"].font.name == "Times New Roman"


def test_document_page_setup_a4():
    """Page size is A4 (210x297mm)."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("Test")
    section = doc.sections[0]
    # Allow 1mm tolerance
    assert abs(section.page_width - Mm(210)) < Mm(1)
    assert abs(section.page_height - Mm(297)) < Mm(1)


def test_document_margins_gost():
    """Margins match GOST R 6.30-2003 (left=30mm, right=15mm, top/bottom=20mm)."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("Test")
    section = doc.sections[0]
    assert abs(section.left_margin - Mm(30)) < Mm(1)
    assert abs(section.right_margin - Mm(15)) < Mm(1)
    assert abs(section.top_margin - Mm(20)) < Mm(1)
    assert abs(section.bottom_margin - Mm(20)) < Mm(1)


def test_heading1_is_centered():
    """H1 style has center alignment."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("# Title")
    h1_style = doc.styles["Heading 1"]
    assert h1_style.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_normal_font_size_12pt():
    """Normal style font is 12pt."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("Test")
    assert doc.styles["Normal"].font.size == Pt(12)


# ── Metadata tests ───────────────────────────────────────────────


def test_document_metadata():
    """Core properties include title and author."""
    converter = MarkdownToDocxConverter(title="My Title", organization_name="My Org")
    doc = converter.convert("# Test")
    assert doc.core_properties.title == "My Title"
    assert doc.core_properties.author == "My Org"
    assert doc.core_properties.language == "ru-RU"


# ── Bytes output ─────────────────────────────────────────────────


def test_convert_to_bytes():
    """convert_to_bytes returns valid DOCX bytes (ZIP/PK signature)."""
    converter = MarkdownToDocxConverter(title="Test")
    result = converter.convert_to_bytes("# Hello")
    assert isinstance(result, bytes)
    assert len(result) > 0
    # DOCX files are ZIP archives starting with PK signature
    assert result[:2] == b"PK"


def test_convert_to_bytes_nonempty():
    """Even minimal text produces a reasonable DOCX."""
    converter = MarkdownToDocxConverter()
    result = converter.convert_to_bytes("Simple text")
    assert len(result) > 1000  # Should be at least a few KB


# ── Empty / edge cases ───────────────────────────────────────────


def test_convert_empty_input():
    """Empty string produces a valid document with no content paragraphs."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("")
    content_paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(content_paras) == 0


def test_convert_only_whitespace():
    """Whitespace-only input produces a valid document."""
    converter = MarkdownToDocxConverter()
    doc = converter.convert("   \n\n   \n")
    content_paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(content_paras) == 0


# ── Full document test ───────────────────────────────────────────


def test_convert_full_legal_document():
    """A realistic full document with all element types renders correctly."""
    md = """# ПОЛИТИКА ОБРАБОТКИ ПЕРСОНАЛЬНЫХ ДАННЫХ

**ООО "Тестовая Компания"**
ИНН: 7701234567

---

## 1. Общие положения

1.1. Настоящая Политика разработана в соответствии с **Федеральным законом от 27.07.2006 № 152-ФЗ** «О персональных данных».

1.2. Оператор обеспечивает обработку персональных данных в *строгом соответствии* с законодательством.

## 2. Категории персональных данных

Оператор обрабатывает следующие категории:

- фамилия, имя, отчество
- адрес электронной почты
- номер телефона

## 3. Цели обработки

| Категория данных | Цель обработки |
|------------------|----------------|
| ФИО | Идентификация |
| Email | Связь |
"""
    converter = MarkdownToDocxConverter(
        title="Политика обработки ПДн",
        organization_name='ООО "Тестовая Компания"',
    )
    doc = converter.convert(md)

    # Verify structure
    h1 = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    h2 = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    bullets = [p for p in doc.paragraphs if p.style.name == "List Bullet"]

    assert len(h1) == 1
    assert "ПОЛИТИКА" in h1[0].text
    assert len(h2) == 3
    assert len(bullets) == 3
    assert len(doc.tables) == 1

    # Check metadata
    assert doc.core_properties.title == "Политика обработки ПДн"


# ── PDF converter tests ──────────────────────────────────────────


def test_pdf_convert_basic():
    """PDF converter produces valid PDF bytes."""
    converter = MarkdownToPdfConverter(title="Тест", organization_name="ООО Тест")
    pdf_bytes = converter.convert_to_bytes("# Заголовок\n\nПараграф текста.")
    assert pdf_bytes[:5] == b"%PDF-"
    assert len(pdf_bytes) > 100


def test_pdf_convert_empty():
    """Empty markdown produces a valid PDF."""
    converter = MarkdownToPdfConverter()
    pdf_bytes = converter.convert_to_bytes("")
    assert pdf_bytes[:5] == b"%PDF-"


def test_pdf_convert_full_document():
    """PDF with headings, lists, bold, table."""
    md = """# Политика обработки ПДн

## 1. Общие положения

Настоящая политика определяет порядок обработки **персональных данных**.

- Пункт первый
- Пункт второй
- Пункт *третий*

1. Нумерованный пункт
2. Второй пункт

| Колонка 1 | Колонка 2 |
|-----------|-----------|
| Значение  | Данные    |

---

## 2. Заключение

Текст заключения.
"""
    converter = MarkdownToPdfConverter(title="Политика", organization_name="ООО Ромашка")
    pdf_bytes = converter.convert_to_bytes(md)
    assert pdf_bytes[:5] == b"%PDF-"
    assert len(pdf_bytes) > 500


def test_create_merged_pdf():
    """Merged PDF contains all documents."""
    docs = [
        {"doc_type": "privacy_policy", "title": "Политика", "content_md": "# Политика\n\nТекст"},
        {"doc_type": "consent_form", "title": "Согласие", "content_md": "# Согласие\n\nТекст"},
    ]
    pdf_bytes = create_merged_pdf(docs, organization_name="Test Org")
    assert pdf_bytes[:5] == b"%PDF-"
    assert len(pdf_bytes) > 500


def test_create_merged_pdf_empty():
    """Empty document list produces a valid PDF."""
    pdf_bytes = create_merged_pdf([])
    assert pdf_bytes[:5] == b"%PDF-"


def test_create_merged_pdf_single_doc():
    """Merged PDF with single document works."""
    docs = [{"doc_type": "cookie_policy", "title": "Cookie", "content_md": "# Cookie\n\nТекст"}]
    pdf_bytes = create_merged_pdf(docs, organization_name="Org")
    assert pdf_bytes[:5] == b"%PDF-"
    assert len(pdf_bytes) > 200
